"""
OPF Privacy Filter Web Application
FastAPI backend for PII detection and redaction.
"""

import asyncio
import json
import logging
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime
from pathlib import Path
from typing import Optional

import csv
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
# OPF_URL removed — detection now uses pf_backend locally
UPLOAD_DIR = Path(tempfile.gettempdir()) / "opf-uploads"
WHITELIST_PATH = Path(os.getenv("WHITELIST_PATH", str(Path(__file__).parent / "whitelist")))
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
OPF_CONCURRENCY = int(os.getenv("OPF_CONCURRENCY", "10"))  # 10 parallel OPF calls
MAX_TASK_CONCURRENCY = int(os.getenv("MAX_TASK_CONCURRENCY", "3"))  # max 3 files processing simultaneously
MIN_SEGMENT_LENGTH = 5  # skip segments shorter than this for OPF detection
TASK_MAX_AGE_HOURS = int(os.getenv("TASK_MAX_AGE_HOURS", "72"))  # auto-clean tasks older than this
TASK_MAX_COUNT = int(os.getenv("TASK_MAX_COUNT", "500"))  # max tasks in memory
_task_sem = asyncio.Semaphore(MAX_TASK_CONCURRENCY)

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
WHITELIST_PATH.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("opf-web")

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------
app = FastAPI(title="OPF Privacy Filter Web", version="1.2.1")

# CORS: allow all origins for local Docker use.
# For production, set CORS_ORIGINS to restrict (e.g., "https://example.com").
_cors_origins = os.getenv("CORS_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=_cors_origins,
    allow_credentials=_cors_origins != ["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files
_static_dir = os.environ.get("STATIC_DIR", "static")
if os.path.isdir(_static_dir):
    app.mount("/static", StaticFiles(directory=_static_dir), name="static")
    _assets_dir = os.path.join(_static_dir, "assets")
    if os.path.isdir(_assets_dir):
        app.mount("/assets", StaticFiles(directory=_assets_dir), name="assets")

# In-memory task store  {task_id: {...}}
tasks: dict[str, dict] = {}
batches: dict[str, dict] = {}  # batch_id -> {batch_id, task_ids, total, created_at}

# ── SQLite persistence for tasks ─────────────────────────────────
import sqlite3

_DB_PATH = Path(os.getenv("TASK_DB", str(WHITELIST_PATH / "tasks.db")))

def _init_db():
    """Create tasks table if it doesn't exist."""
    conn = sqlite3.connect(str(_DB_PATH))
    conn.execute("""
        CREATE TABLE IF NOT EXISTS tasks (
            task_id    TEXT PRIMARY KEY,
            filename   TEXT,
            filepath   TEXT,
            ext        TEXT,
            status     TEXT,
            progress   REAL DEFAULT 0,
            error      TEXT,
            created_at TEXT,
            result_json TEXT,
            parsed_sections_json TEXT
        )
    """)
    conn.commit()
    conn.close()

def _save_task(task_id: str):
    """Flush one task to SQLite."""
    t = tasks.get(task_id)
    if not t:
        return
    try:
        conn = sqlite3.connect(str(_DB_PATH))
        result = t.get("result")
        parsed = t.get("parsed_sections")
        conn.execute(
            """INSERT OR REPLACE INTO tasks
               (task_id, filename, filepath, ext, status, progress, error, created_at, result_json, parsed_sections_json)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                task_id,
                t.get("filename"),
                t.get("filepath"),
                t.get("ext"),
                t.get("status"),
                t.get("progress", 0),
                t.get("error"),
                t.get("created_at"),
                json.dumps(result, ensure_ascii=False) if result else None,
                json.dumps(parsed, ensure_ascii=False) if parsed else None,
            ),
        )
        conn.commit()
        conn.close()
    except Exception as e:
        logger.warning("Failed to save task %s to DB: %s", task_id, e)

def _load_tasks():
    """Load all tasks from SQLite into memory on startup."""
    if not _DB_PATH.exists():
        return
    try:
        conn = sqlite3.connect(str(_DB_PATH))
        conn.row_factory = sqlite3.Row
        rows = conn.execute("SELECT * FROM tasks").fetchall()
        conn.close()
        count = 0
        for row in rows:
            tid = row["task_id"]
            tasks[tid] = {
                "task_id": tid,
                "filename": row["filename"],
                "filepath": row["filepath"],
                "ext": row["ext"],
                "status": row["status"],
                "progress": row["progress"] or 0,
                "error": row["error"],
                "created_at": row["created_at"],
                "result": json.loads(row["result_json"]) if row["result_json"] else {},
                "parsed_sections": json.loads(row["parsed_sections_json"]) if row["parsed_sections_json"] else [],
            }
            count += 1
        if count:
            logger.info("Loaded %d tasks from SQLite", count)
    except Exception as e:
        logger.warning("Failed to load tasks from DB: %s", e)

def _delete_task_from_db(task_id: str):
    """Remove one task from SQLite."""
    try:
        conn = sqlite3.connect(str(_DB_PATH))
        conn.execute("DELETE FROM tasks WHERE task_id = ?", (task_id,))
        conn.commit()
        conn.close()
    except Exception:
        pass

def _clear_all_tasks_db():
    """Wipe all tasks from SQLite."""
    try:
        conn = sqlite3.connect(str(_DB_PATH))
        conn.execute("DELETE FROM tasks")
        conn.commit()
        conn.close()
    except Exception:
        pass

# Initialize DB and load existing tasks
_init_db()
_load_tasks()

# Mark interrupted tasks (server restart killed their processing coroutine)
for tid, t in list(tasks.items()):
    if t.get("status") in ("queued", "processing", "parsing", "detecting"):
        t["status"] = "error"
        t["error"] = "服务重启，任务中断 — 请点击重新扫描"
        _save_task(tid)
        logger.info("Marked interrupted task %s as error", tid)

# ---------------------------------------------------------------------------
# File parsers
# ---------------------------------------------------------------------------

def parse_text(filepath: Path) -> list[dict]:
    text = filepath.read_text(encoding="utf-8", errors="replace")
    return [{"type": "paragraph", "index": 0, "text": text}]


def parse_csv(filepath: Path) -> list[dict]:
    sections: list[dict] = []
    with open(filepath, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        rows = [row for row in reader]
    sections.append({"type": "table", "index": 0, "rows": rows})
    return sections


def parse_xlsx(filepath: Path) -> list[dict]:
    import openpyxl

    sections: list[dict] = []
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    for idx, ws in enumerate(wb.worksheets):
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append([str(cell) if cell is not None else "" for cell in row])
        sections.append({"type": "table", "index": idx, "rows": rows, "sheet": ws.title})
    wb.close()
    return sections


def parse_docx(filepath: Path) -> list[dict]:
    from docx import Document

    doc = Document(str(filepath))
    sections: list[dict] = []
    idx = 0
    for para in doc.paragraphs:
        if para.text.strip():
            sections.append({"type": "paragraph", "index": idx, "text": para.text})
            idx += 1
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        sections.append({"type": "table", "index": idx, "rows": rows})
        idx += 1
    return sections


def parse_image(filepath: Path) -> list[dict]:
    text = ocr_extract_text(str(filepath))
    if text.strip():
        return [{"type": "paragraph", "index": 0, "text": text}]
    return []


def parse_pdf(filepath: Path) -> list[dict]:
    import pdfplumber

    sections: list[dict] = []
    ocr_candidates = []  # pages to OCR: no text OR has images

    with pdfplumber.open(filepath) as pdf:
        # Phase 1: Quick scan — does ANY page need OCR?
        any_images = False
        for page in pdf.pages:
            if getattr(page, 'images', None):
                any_images = True
                break

        # Phase 2: Extract text + tables, collect OCR candidates
        for page_idx, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            has_text = bool(text.strip())
            has_images = bool(getattr(page, 'images', None)) if any_images else False

            if has_text:
                sections.append({"type": "paragraph", "index": page_idx, "text": text})

            # OCR only if: page has images (scanned doc, ID photo, etc.)
            # or page has no extractable text at all
            if has_images or not has_text:
                ocr_candidates.append((page_idx, page))

            tables = page.extract_tables()
            for tbl_idx, table in enumerate(tables):
                rows = [[str(cell) if cell else "" for cell in row] for row in table]
                sections.append({"type": "table", "index": page_idx, "rows": rows, "page": page_idx + 1})

        # Phase 3: OCR only if needed — skip entirely if no images and all pages have text
        if ocr_candidates and (not any_images and all(s.get("text") for s in sections if s["type"] == "paragraph")):
            logger.info("PDF has text on all pages and no images — skipping OCR")
            ocr_candidates = []

        # Batch OCR for large PDFs — process in chunks to control memory
        OCR_BATCH_SIZE = 30  # pages per batch
        total_ocr = len(ocr_candidates)

        if total_ocr > OCR_BATCH_SIZE:
            logger.info("Large PDF: %d OCR candidates, processing in batches of %d",
                        total_ocr, OCR_BATCH_SIZE)

        for batch_start in range(0, total_ocr, OCR_BATCH_SIZE):
            batch = ocr_candidates[batch_start:batch_start + OCR_BATCH_SIZE]
            batch_num = batch_start // OCR_BATCH_SIZE + 1
            total_batches = (total_ocr + OCR_BATCH_SIZE - 1) // OCR_BATCH_SIZE

            if total_ocr > OCR_BATCH_SIZE:
                logger.info("OCR batch %d/%d: pages %d-%d",
                            batch_num, total_batches,
                            batch[0][0], batch[-1][0])

            for page_idx, page in batch:
                try:
                    # Skip pages with tiny images (likely stamps/seals, no PII)
                    images = getattr(page, 'images', []) or []
                    if images and len(images) <= 2:
                        total_img_area = sum(
                            (img.get('x1', 0) - img.get('x0', 0)) * (img.get('y1', 0) - img.get('y0', 0))
                            for img in images
                        )
                        page_area = page.width * page.height
                        if page_area > 0 and total_img_area / page_area < 0.15:
                            logger.debug("Page %d: stamp page (%.0f%% of page), skipping OCR",
                                         page_idx, total_img_area / page_area * 100)
                            continue

                    # ── Dynamic OCR resolution ──
                    # Strategy: if page has images but very little text (watermark/overlay only),
                    # it's likely an ID card, certificate, or receipt → boost to 300dpi
                    # Full-page text scans with images → 150dpi is sufficient
                    dpi = 150
                    if images:
                        # Check if the extracted text is just watermarks/overlays (< 100 chars)
                        page_text = (page.extract_text() or '').strip()
                        text_len = len(page_text)
                        page_area = page.width * page.height
                        total_img_area = sum(
                            (img.get('x1', 0) - img.get('x0', 0)) * (img.get('y1', 0) - img.get('y0', 0))
                            for img in images
                        )
                        img_ratio = total_img_area / page_area if page_area > 0 else 0

                        if text_len < 100 and img_ratio > 0.3:
                            # Small text + big image = ID card / certificate / receipt → high DPI
                            dpi = 300
                            logger.info("Page %d: image-heavy, minimal text (%d chars, %.0f%% img) → 300dpi",
                                        page_idx, text_len, img_ratio * 100)
                        elif img_ratio < 0.3:
                            # Small embedded image → likely needs high DPI to read
                            dpi = 300
                            logger.info("Page %d: small image (%.0f%% of page) → 300dpi",
                                        page_idx, img_ratio * 100)

                    img = page.to_image(resolution=dpi)
                    img_path = str(filepath.parent / f"_ocr_temp_{page_idx}.png")
                    img.save(img_path, format="PNG")
                    try:
                        ocr_text = ocr_extract_text(img_path)
                        if ocr_text.strip():
                            sections.append({
                                "type": "paragraph",
                                "index": page_idx,
                                "text": ocr_text,
                                "source": "ocr",
                            })
                    finally:
                        Path(img_path).unlink(missing_ok=True)
                except Exception as e:
                    logger.warning("OCR fallback failed for page %d: %s", page_idx, e)

            # Release batch memory explicitly
            del batch
            import gc
            gc.collect()

    return sections


PARSERS = {
    ".txt": parse_text,
    ".md": parse_text,
    ".csv": parse_csv,
    ".xlsx": parse_xlsx,
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".jpg": parse_image,
    ".jpeg": parse_image,
    ".png": parse_image,
    ".bmp": parse_image,
    ".tiff": parse_image,
    ".tif": parse_image,
}

# ---------------------------------------------------------------------------
# File type detection via Magika (Google AI-powered, offline)
# ---------------------------------------------------------------------------
_magika_instance = None

def get_magika():
    global _magika_instance
    if _magika_instance is None:
        from magika import Magika
        _magika_instance = Magika()
    return _magika_instance


# Magika label → parser extension mapping
MAGIKA_LABEL_TO_EXT: dict[str, str] = {
    # Documents
    "pdf": ".pdf",
    "docx": ".docx",
    "doc": ".docx",
    "rtf": ".docx",
    "xlsx": ".xlsx",
    "xls": ".xlsx",
    "ods": ".xlsx",
    "csv": ".csv",
    "tsv": ".csv",
    # Markdown
    "markdown": ".md",
    # Text / code → treat as plain text
    "txt": ".txt", "python": ".txt", "javascript": ".txt", "typescript": ".txt",
    "json": ".txt", "xml": ".txt", "html": ".txt", "css": ".txt",
    "java": ".txt", "c": ".txt", "cpp": ".txt", "csharp": ".txt",
    "go": ".txt", "rust": ".txt", "ruby": ".txt", "php": ".txt",
    "perl": ".txt", "lua": ".txt", "r": ".txt", "scala": ".txt",
    "swift": ".txt", "kotlin": ".txt", "shell": ".txt", "bash": ".txt",
    "yaml": ".txt", "toml": ".txt", "ini": ".txt", "log": ".txt",
    "sql": ".txt", "makefile": ".txt", "dockerfile": ".txt",
    # Tcl / Tickless labels sometimes appear for small text snippets
    "tcl": ".txt", "tickle": ".txt",
    # Images (for OCR)
    "jpeg": ".jpg",
    "png": ".png",
    "bmp": ".bmp",
    "tiff": ".tiff",
}


def detect_content_type(content: bytes, filename: str) -> tuple[str | None, str, str]:
    """Detect file content type using Magika.
    Returns (parser_ext_or_none, label, description).
    """
    try:
        magika = get_magika()
        result = magika.identify_bytes(content)
        label = result.output.label
        description = result.output.description
        ext = MAGIKA_LABEL_TO_EXT.get(label)
        return ext, label, description
    except Exception as e:
        logger.warning("Magika detection failed: %s, falling back to extension", e)
        fallback_ext = Path(filename).suffix.lower()
        return fallback_ext, "unknown", "unknown"


# ---------------------------------------------------------------------------
# Whitelist helpers
# ---------------------------------------------------------------------------

def load_whitelist() -> dict:
    wl_file = WHITELIST_PATH / "pii_whitelist.json"
    if wl_file.exists():
        return json.loads(wl_file.read_text(encoding="utf-8"))
    return {"version": "1.2", "rules": []}


def save_whitelist(data: dict) -> None:
    wl_file = WHITELIST_PATH / "pii_whitelist.json"
    wl_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


# ---------------------------------------------------------------------------
# File scan cache (hash-based, skip re-detection for identical content)
# ---------------------------------------------------------------------------

def _file_hash(content: bytes) -> str:
    import hashlib
    return hashlib.sha256(content).hexdigest()[:16]


CACHE_VERSION = "8.0"  # Bump when result schema changes


def load_scan_cache() -> dict:
    cache_file = WHITELIST_PATH / "scan_cache.json"
    if cache_file.exists():
        data = json.loads(cache_file.read_text(encoding="utf-8"))
        # Auto-invalidate if cache version doesn't match
        if data.get("version") != CACHE_VERSION:
            return {"version": CACHE_VERSION, "entries": {}}
        return data
    return {"version": CACHE_VERSION, "entries": {}}


def save_scan_cache(data: dict) -> None:
    data["version"] = CACHE_VERSION
    cache_file = WHITELIST_PATH / "scan_cache.json"
    cache_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def cache_get(file_hash: str) -> dict | None:
    cache = load_scan_cache()
    return cache["entries"].get(file_hash)


def cache_put(file_hash: str, result: dict) -> None:
    cache = load_scan_cache()
    cache["entries"][file_hash] = {
        "total_pii": result.get("total_pii", 0),
        "text_segments": result.get("text_segments", 0),
        "char_count": result.get("char_count", 0),
        "line_count": result.get("line_count", 0),
        "by_type": result.get("by_type", {}),
        "details": result.get("details", []),
        "cached_at": datetime.now().isoformat(),
    }
    # Trim cache to 500 entries (LRU-ish: remove oldest)
    if len(cache["entries"]) > 500:
        sorted_keys = sorted(cache["entries"], key=lambda k: cache["entries"][k].get("cached_at", ""))
        for k in sorted_keys[:len(cache["entries"]) - 500]:
            del cache["entries"][k]
    save_scan_cache(cache)


# ---------------------------------------------------------------------------
# Regex safety check (ReDoS prevention)
# ---------------------------------------------------------------------------
_RE_DOS_TIMEOUT = 5  # seconds

def is_safe_regex(pattern: str) -> tuple[bool, str]:
    """Validate a regex pattern for safety (no catastrophic backtracking).

    Tests compilation + short match on a crafted input.
    Returns (is_safe, error_message).
    """
    import re as _re

    try:
        compiled = _re.compile(pattern)
    except _re.error as e:
        return False, f"Regex 编译错误: {e}"

    # ReDoS probe: use a crafted input designed to trigger catastrophic backtracking
    # For patterns with nested quantifiers like (a+)+b, feed "aaaaaa" and see if it times out
    probe = "a" * 30

    def _timed_match():
        try:
            compiled.search(probe)
            return True, ""
        except Exception as e:
            return False, str(e)

    import threading
    result = [True, ""]

    def _run():
        safe, msg = _timed_match()
        result[0] = safe
        result[1] = msg

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    t.join(timeout=_RE_DOS_TIMEOUT)

    if t.is_alive():
        return False, "正则表达式存在 ReDoS 风险（触发灾难性回溯）"

    if not result[0]:
        return False, f"正则表达式匹配失败: {result[1]}"

    return True, ""


# ---------------------------------------------------------------------------
# Custom sensitive word dictionary (user-defined PII patterns)
# ---------------------------------------------------------------------------

def load_dictionary() -> dict:
    dict_file = WHITELIST_PATH / "pii_dictionary.json"
    if dict_file.exists():
        return json.loads(dict_file.read_text(encoding="utf-8"))
    return {"version": "1.0", "words": []}


def save_dictionary(data: dict) -> None:
    dict_file = WHITELIST_PATH / "pii_dictionary.json"
    dict_file.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def match_dictionary(text: str) -> list[dict]:
    """Find all custom dictionary words in text. Returns list of spans."""
    dictionary = load_dictionary()
    spans = []
    for entry in dictionary.get("words", []):
        pattern = entry.get("text", "")
        label = entry.get("type", "custom_pii")
        if not pattern:
            continue
        # Find all occurrences
        start = 0
        while True:
            idx = text.find(pattern, start)
            if idx < 0:
                break
            spans.append({
                "label": label,
                "start": idx,
                "end": idx + len(pattern),
                "text": pattern,
            })
            start = idx + len(pattern)
    return spans

# ---------------------------------------------------------------------------
# Built-in PII regex patterns (Chinese-specific, supplements OPF)
# ---------------------------------------------------------------------------

# ── 中国手机号 ──────────────────────────────────────────────────
# 运营商号段（工信部 2026 年最新分配）:
#   移动: 134-139, 147, 148, 150-152, 157-159, 165, 172, 178, 182-184, 187-188, 195-198
#   联通: 130-132, 145, 146, 155, 156, 166, 167, 171, 175, 176, 185, 186, 196
#   电信: 133, 149, 153, 173, 174, 177, 180, 181, 189, 190, 191, 193, 199
#   广电: 192
#   虚拟运营商: 162, 165, 167, 170, 171
# 简化正则: 1[3-9]X XXXX XXXX（覆盖全部已分配号段，排除未分配的 10x/11x/12x）
_PHONE_RE = re.compile(r'(?<!\d)(1[3-9]\d{9})(?!\d)')

# ── 中国大陆座机号 ──────────────────────────────────────────────
# 区号(3-4位) + 号码(7-8位)，可选分隔符 - 或 空格
# 区号: 010, 020-029 (两位区号), 03xx-09xx (三位区号)
_LANDLINE_RE = re.compile(
    r'(?<!\d)'
    r'(0(?:10|2[0-9]|[3-9]\d{1,2})[-\s]?\d{7,8})'
    r'(?!\d)'
)

# ── 400/800 企业热线 ─────────────────────────────────────────────
_400_800_RE = re.compile(r'(?<!\d)((?:400|800)[-\s]?\d{3,4}[-\s]?\d{4})(?!\d)')

# ── 银行卡号 ────────────────────────────────────────────────────
_BANKCARD_RE = re.compile(r'(?<!\d)(\d{16,20})(?!\d)')

# ── 身份证号（18位，含校验位X）──────────────────────────────────
# 6位地区码 + 8位出生日期 + 3位顺序码 + 1位校验码
_IDCARD_RE = re.compile(
    r'(?<!\d)'
    r'([1-9]\d{5}'                           # 地区码（非0开头）
    r'(?:19|20)\d{2}'                        # 年份 1900-2099
    r'(?:0[1-9]|1[0-2])'                     # 月份 01-12
    r'(?:0[1-9]|[12]\d|3[01])'              # 日期 01-31
    r'\d{3}'                                 # 顺序码
    r'[\dXx])'                               # 校验码
    r'(?!\d)'
)

# ── 公网 IP 地址 ────────────────────────────────────────────────
_IP_RE = re.compile(
    r'(?<!\d)'
    r'((?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]\d|\d)\.'
    r'(?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]\d|\d)\.'
    r'(?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]\d|\d)\.'
    r'(?:25[0-5]|2[0-4]\d|1\d{2}|[1-9]\d|\d))'
    r'(?!\d)'
)

# ── 大模型 / SaaS API Key ──────────────────────────────────────
# OpenAI/DeepSeek/Kimi/MiMo/SiliconFlow: sk-[a-zA-Z0-9_-]{20,}
# Anthropic/Claude: sk-ant-api03-... (sk-ant- prefix)
# Azure OpenAI: 32-hex key (8a2b3c4d...)
# Google AI: AIzaSy... (39 chars)
# HuggingFace: hf_[a-zA-Z0-9]{20,}
# 阿里云 AccessKey: LTAI[a-zA-Z0-9]{12,}
# AWS AccessKey: AKIA[A-Z0-9]{16}
_APIKEY_RE = re.compile(
    r'(sk-[a-zA-Z0-9_-]{20,})'              # OpenAI/DeepSeek/Kimi/MiMo/SiliconFlow/Claude
    r'|(sk-ant-[a-zA-Z0-9_-]{20,})'         # Anthropic Claude
    r'|(hf_[a-zA-Z0-9]{20,})'               # HuggingFace
    r'|(AIzaSy[a-zA-Z0-9_-]{30,})'          # Google AI
    r'|(LTAI[a-zA-Z0-9]{12,})'              # 阿里云 AccessKey ID
    r'|(AKIA[A-Z0-9]{16})'                  # AWS AccessKey ID
)

# ── 中国身份证号校验 ────────────────────────────────────────────
def _validate_idcard(num: str) -> bool:
    """Validate Chinese ID card number with checksum."""
    if len(num) != 18:
        return False
    weights = [7, 9, 10, 5, 8, 4, 2, 1, 6, 3, 7, 9, 10, 5, 8, 4, 2]
    check_codes = '10X98765432'
    try:
        total = sum(int(num[i]) * weights[i] for i in range(17))
        return check_codes[total % 11] == num[17].upper()
    except (ValueError, IndexError):
        return False


def detect_builtin_patterns(text: str) -> list[dict]:
    """Detect PII via regex: phone, landline, ID card, bank card, IP, API key.
    Supplements OPF model detection for better Chinese coverage.
    """
    spans = []

    # ── 手机号 ──
    for m in _PHONE_RE.finditer(text):
        spans.append({"label": "private_phone", "start": m.start(), "end": m.end(), "text": m.group()})

    # ── 座机号 ──
    for m in _LANDLINE_RE.finditer(text):
        spans.append({"label": "private_phone", "start": m.start(), "end": m.end(), "text": m.group()})

    # ── 400/800 企业热线 ──
    for m in _400_800_RE.finditer(text):
        spans.append({"label": "private_phone", "start": m.start(), "end": m.end(), "text": m.group()})

    # ── 身份证号（带校验）──
    for m in _IDCARD_RE.finditer(text):
        num = m.group()
        if _validate_idcard(num):
            spans.append({"label": "private_idcard", "start": m.start(), "end": m.end(), "text": num})

    # ── 银行卡号 ──
    for m in _BANKCARD_RE.finditer(text):
        num = m.group()
        # Skip if it looks like an ID card (18 digits starting with valid area code)
        if len(num) == 18 and num[0] in '123456' and re.match(r'^\d{17}[\dXx]$', num):
            continue  # likely ID card, not bank card
        if 16 <= len(num) <= 20:
            known_prefixes = ('62', '60', '56', '45', '44', '43', '42', '41', '40',
                              '37', '36', '35', '34', '30', '19', '17', '11', '10',
                              '98', '95', '69')
            if any(num.startswith(p) for p in known_prefixes) or _luhn_check(num):
                spans.append({"label": "private_bankcard", "start": m.start(), "end": m.end(), "text": num})

    # ── 公网 IP 地址（排除私网）──
    for m in _IP_RE.finditer(text):
        ip = m.group()
        parts = ip.split('.')
        if len(parts) != 4:
            continue
        first, second = int(parts[0]), int(parts[1])
        if first == 10:
            continue
        if first == 172 and 16 <= second <= 31:
            continue
        if first == 192 and second == 168:
            continue
        if first in (127, 0, 169) or first >= 224:
            continue
        spans.append({"label": "private_url", "start": m.start(), "end": m.end(), "text": ip})

    # ── 大模型 API Key（OpenAI/DeepSeek/Claude/Kimi/MiMo 等）──
    for m in _APIKEY_RE.finditer(text):
        spans.append({"label": "secret", "start": m.start(), "end": m.end(), "text": m.group()})

    return spans


def _luhn_check(num: str) -> bool:
    """Luhn algorithm — validates credit card / bank card numbers."""
    try:
        digits = [int(d) for d in num]
        odd_digits = digits[-1::-2]
        even_digits = digits[-2::-2]
        checksum = sum(odd_digits)
        for d in even_digits:
            checksum += sum(divmod(d * 2, 10))
        return checksum % 10 == 0
    except (ValueError, IndexError):
        return False


# ---------------------------------------------------------------------------
# OnnxOCR for scanned PDFs and image files (lazy init, optional)
# Lightweight PP-OCRv5 ONNX-based OCR, ~126MB vs PaddleOCR's ~581MB
# ---------------------------------------------------------------------------
OCR_ENABLED = os.getenv("OCR_ENABLED", "true").lower() in ("true", "1", "yes")
_ocr_instance = None
_ocr_unavailable = False  # Set True if OCR fails to init (don't retry)

def _check_memory():
    """Check if enough free memory for OCR (~300MB needed for OnnxOCR)."""
    try:
        import psutil
        avail_mb = psutil.virtual_memory().available / (1024 * 1024)
        return avail_mb > 500  # OnnxOCR needs ~300MB, keep 200MB margin
    except Exception:
        return True  # Can't check, allow

def get_ocr():
    global _ocr_instance, _ocr_unavailable
    if _ocr_unavailable:
        return None
    if _ocr_instance is None:
        if not OCR_ENABLED:
            logger.info("OCR disabled via OCR_ENABLED=false")
            _ocr_unavailable = True
            return None
        if not _check_memory():
            logger.warning("Insufficient memory for OCR (need ~500MB free), skipping")
            _ocr_unavailable = True
            return None
        try:
            from onnxocr.onnx_paddleocr import ONNXPaddleOcr
            _ocr_instance = ONNXPaddleOcr(use_angle_cls=False, use_gpu=False, det_limit_side_len=640)
            logger.info("OnnxOCR initialized successfully")
        except Exception as e:
            logger.error("Failed to initialize OnnxOCR: %s", e)
            _ocr_unavailable = True
            return None
    return _ocr_instance

def ocr_extract_text(image_path: str) -> str:
    """Extract text from image using OnnxOCR (PP-OCRv5 ONNX)."""
    ocr = get_ocr()
    if ocr is None:
        logger.warning("OCR unavailable, skipping image: %s", image_path)
        return ""
    texts = []
    try:
        import cv2
        img = cv2.imread(image_path)
        if img is None:
            logger.warning("Failed to read image: %s", image_path)
            return ""
        result = ocr.ocr(img, cls=False)
        if result and result[0]:
            for line in result[0]:
                text = line[1][0] if isinstance(line[1], (list, tuple)) else str(line[1])
                conf = line[1][1] if isinstance(line[1], (list, tuple)) else 0
                if text and text.strip() and conf > 0.5:
                    texts.append(text.strip())
    except Exception as e:
        logger.error("OCR inference failed for %s: %s", image_path, e)
    return '\n'.join(texts)


def merge_pii_spans(opf_spans: list[dict], ner_spans: list[dict]) -> list[dict]:
    """Merge OPF and HanLP spans, deduplicate overlaps (OPF takes priority)."""
    if not ner_spans:
        return opf_spans
    if not opf_spans:
        return ner_spans

    # Build set of character ranges covered by OPF
    opf_chars = set()
    for s in opf_spans:
        opf_chars.update(range(s["start"], s["end"]))

    # Add NER spans that don't overlap with OPF
    merged = list(opf_spans)
    for ns in ner_spans:
        ner_range = set(range(ns["start"], ns["end"]))
        if not ner_range & opf_chars:
            merged.append(ns)

    merged.sort(key=lambda x: x["start"])
    return merged


# ---------------------------------------------------------------------------
# PII detection via pf_backend (privacy-filter.cpp)
# ---------------------------------------------------------------------------

# Lazy singleton — initialised on first call so imports don't require the model
_pf_instance = None

def _get_pf():
    global _pf_instance
    if _pf_instance is None:
        from pf_backend import PFBackend
        _pf_instance = PFBackend.get()
    return _pf_instance


async def detect_pii_batch(texts: list[str]) -> list[dict]:
    """Detect PII in a batch of texts using pf_backend (local ctypes call).

    Returns a list of dicts with "detected_spans" key to match the old
    OPF /redact/batch response format consumed by the caller.
    """
    pf = _get_pf()
    loop = asyncio.get_running_loop()

    def _classify_one(text: str) -> dict:
        spans = pf.classify(text, threshold=0.5)
        return {"detected_spans": spans}

    # Run blocking classify calls in a thread pool
    results = await asyncio.gather(
        *[loop.run_in_executor(None, _classify_one, t) for t in texts]
    )
    return list(results)


def _post_process_spans(text: str, opf_spans: list[dict]) -> list[dict]:
    """Merge OPF spans with dict + builtin, then filter false positives.

    Note: OPF model + regex engine covers all detection needs.
    """
    dict_spans = match_dictionary(text)
    builtin_spans = detect_builtin_patterns(text)
    spans = merge_pii_spans(opf_spans, dict_spans)
    spans = merge_pii_spans(spans, builtin_spans)
    spans = [s for s in spans if not _is_false_positive(s["text"], s["label"])]
    # Label-aware space filter: only strip spaces for types where spaces
    # are genuinely suspicious (phone, ID card, API key, bank card, pure number sequences).
    # Keep spans for addresses, names, organizations, emails, URLs that may contain spaces.
    _space_sensitive_labels = {
        'private_phone', 'PHONE', 'PRIVATE_PHONE',
        'private_idcard', 'IDCARD', 'PRIVATE_IDCARD',
        'account_number',  # pf_backend: ID card / SSN / account numbers
        'private_bankcard', 'BANK', 'PRIVATE_BANKCARD',
        'secret', 'SECRET',
    }
    spans = [
        s for s in spans
        if s["label"] not in _space_sensitive_labels or ' ' not in s.get("text", "")
    ]
    return spans


def _is_false_positive(detected_text: str, label: str) -> bool:
    """Check if a detected PII span is likely a false positive.
    Filters: dates, pure numbers, currency amounts, version numbers, etc.
    IMPORTANT: Never filter out spans that OPF already classified as real PII.
    """
    t = detected_text.strip()

    # Filter private/reserve IP addresses (OPF model returns these as private_url)
    if label in ('private_url', 'URL', 'PRIVATE_URL'):
        ip_match = re.match(r'^(\d{1,3})\.(\d{1,3})\.(\d{1,3})\.(\d{1,3})$', t)
        if ip_match:
            first, second = int(ip_match.group(1)), int(ip_match.group(2))
            if first == 10 or (first == 172 and 16 <= second <= 31) or (first == 192 and second == 168) or first in (127, 0, 169) or first >= 224:
                return True  # private/reserved IP, not PII

    # OPF already classified these — generally trust, but still validate
    real_pii_labels = {
        'private_phone', 'PHONE', 'PRIVATE_PHONE',
        'private_email', 'EMAIL', 'PRIVATE_EMAIL',
        'private_idcard', 'IDCARD', 'PRIVATE_IDCARD',
        'account_number',  # pf_backend: ID card / SSN / account numbers
        'private_bankcard', 'BANK', 'PRIVATE_BANKCARD',
        'secret', 'SECRET',
    }
    if label in real_pii_labels:
        # Phone-specific validation: OPF sometimes over-detects short numbers as phone
        if label in ('private_phone', 'PHONE', 'PRIVATE_PHONE'):
            digits = re.sub(r'\D', '', t)
            digit_count = len(digits)
            # Too short to be any phone number (minimum: 7 digits international)
            if digit_count < 7:
                return True
            # Chinese mobile: must be 11 digits starting with 1[3-9]
            if digit_count == 11 and not digits.startswith(('13', '14', '15', '16', '17', '18', '19')):
                return True
            # Not a phone length at all (7, 8, 10, 11 are valid; everything else is suspicious)
            if digit_count not in (7, 8, 9, 10, 11, 12, 13, 14, 15):
                return True
            # Pure float/decimal like 33.0, 1.5 — not a phone
            if re.match(r'^[\d,]+\.?\d*$', t) and '.' in t:
                return True
        return False

    # Pure numbers (integers, floats, with optional commas, with optional minus)
    # BUT exclude phone-length (11) and bank-account-length (16-20) sequences
    if re.match(r'^-?[\d,]+\.?\d*$', t):
        digits_only = re.sub(r'[,.\-]', '', t)
        digit_count = len(digits_only)
        # 11 digits = likely phone; 16-20 digits = likely bank account
        if digit_count == 11 and digits_only.startswith('1'):
            return False  # could be phone
        if 16 <= digit_count <= 20:
            return False  # could be bank account
        return True

    # Date patterns: 2025-12, 2026-01-01, 2025/12, 20260101, 202512
    if re.match(r'^\d{4}[-/]\d{1,2}([-/]\d{1,2})?$', t):
        return True
    if re.match(r'^\d{4}年\d{1,2}月(\d{1,2}日)?$', t):
        return True
    if re.match(r'^\d{6,8}$', t) and not re.match(r'^1[3-9]\d{9}$', t):
        return True

    # Negative amounts: -31.65, -1234.56
    if re.match(r'^-[\d,]+\.?\d*$', t):
        return True

    # Currency/amount patterns: ¥1000, $500, 1000元, 3.14万
    if re.match(r'^[¥$€£￥][\d,]+\.?\d*$', t):
        return True
    if re.match(r'^[\d,]+\.?\d*[元万亿]$', t):
        return True

    # Version numbers: v1.2.3, 1.0.0, 3.12-slim
    if re.match(r'^v?\d+\.\d+(\.\d+)?(-\w+)?$', t):
        return True

    # IP-like but actually version/float: 3.12, 2.6.1
    if re.match(r'^\d+\.\d+(\.\d+)?$', t) and len(t) < 10:
        return True

    # Pure percentage: 45%, 0.5%
    if re.match(r'^[\d.]+%$', t):
        return True

    # Cloud resource instance IDs (eip-, i-, ecs-, slb-, rds-, nas-, vpc-, sg-...)
    if re.match(r'^(eip|i|ecs|slb|rds|nas|vpc|sg|sg_|eni|nat|cdn|oss|slb|mse|hbase|drds|kvstore|polardb|elasticsearch|log|cms|ess|ros|fnf|oos|kms|cr|cs|csb|swas|ga|vpc|cen|tr|ccn|dc|sag|cfw|waf|avds|havip|eipanycast|vpclattice|nlb|alb|gwlb|clb)[-\w]*$', t, re.IGNORECASE):
        return True

    # Common Chinese words mis-tagged as person names
    # Categories: tech/platforms, security, operations, business, general vocabulary
    false_person_words = {
        # 科技平台
        '博客', '微博', '微信', '抖音', '快手', '百度', '淘宝', '京东',
        '美团', '滴滴', '字节', '网易', '新浪', '搜狐', '谷歌', '微软',
        '苹果', '亚马逊', '腾讯', '阿里', '华为', '小米', 'OPPO',
        # 安全术语
        '白名单', '黑名单', '安全法', '安全性', '安全防护', '安全策略',
        '安全审计', '安全检查', '安全事件', '安全意识', '安全保障',
        '信息安全', '网络安全', '通信安全', '数据安全', '加密传输',
        '加密协议', '加密算法', '访问控制', '身份认证', '权限管理',
        '漏洞扫描', '入侵检测', '防火墙', '安全网关', '安全基线',
        '安全等级', '等保', '合规', '合规性', '风险评估', '安全评估',
        # 文档术语（常被 OPF 误判为人名）
        '承诺书', '申请人', '盖章处', '授权人', '审批人', '经办人',
        '联系人', '收件人', '发件人', '抄送人', '备注',
        # 技术术语
        '智能化', '自动化', '数字化', '信息化', '云化', '虚拟化',
        '终端设备', '网络设备', '服务器', '交换机', '路由器',
        '严格遵守', '严格执行', '严格管理', '严格落实',
        '范畴', '卓越', '全面提升', '有效保障', '切实',
        '互联互通', '资源共享', '协同', '一体化',
        # 运维术语
        '触达', '信媒', '财务', '国际', '短信', '平台', '资源',
        '总成本', '合计', '默认', '测试', '生产', '开发', '运维',
        '日志', '监控', '告警', '备份', '存储', '网络', '安全',
        '数据库', '中间件', '容器', '集群', '节点', '服务',
        '系统', '模块', '组件', '配置', '管理', '控制', '中心',
        '引擎', '网关', '代理', '缓存', '队列', '管道', '接口',
        '协议', '标准', '规范', '架构', '方案', '流程', '策略',
        '机制', '体系', '制度', '规范', '要求', '措施', '办法',
        # 通用动词/形容词（常被误判）
        '加强', '完善', '提升', '优化', '推进', '落实', '确保',
        '保障', '维护', '建设', '实施', '开展', '推动', '促进',
        '规范', '健全', '强化', '深化', '细化', '明确', '落实',
    }
    if label in ('private_person', 'PERSON') and t in false_person_words:
        return True

    # Also filter 2-char non-name words that OPF mislabels
    # If a "person name" is really a common word (not a surname+given name), filter it
    if label in ('private_person', 'PERSON') and len(t) <= 3:
        # These 2-char combos are almost never real names in business docs
        _non_name_suffixes = {
            '安全', '保障', '管理', '系统', '网络', '通信', '数据', '信息',
            '技术', '服务', '设备', '平台', '中心', '方案', '策略', '制度',
            '规范', '标准', '要求', '措施', '流程', '架构', '体系', '机制',
            '风险', '评估', '检查', '审计', '监控', '防护', '传输', '加密',
            '控制', '认证', '权限', '漏洞', '备份', '存储', '日志', '告警',
        }
        if t in _non_name_suffixes:
            return True

    # Common Chinese words mis-tagged as address
    false_address_words = {
        '服务', '日志', '监控', '平台', '系统', '模块', '组件',
        '配置', '管理', '控制', '中心', '引擎', '网关', '代理',
        '缓存', '队列', '管道', '接口', '协议', '标准', '规范',
    }
    if label in ('private_address', 'ADDRESS') and t in false_address_words:
        return True

    # Common words mis-tagged as organization
    false_org_words = {
        '网络通信', '加密协议', '安全防护', '信息安全', '数据安全',
        '网络安全', '通信安全', '访问控制', '身份认证',
        '中国移动', '中国联通', '中国电信', '中国广电',  # 运营商名称，非个人所属
        '阿里云', '腾讯云', '华为云', '百度云', '亚马逊云',
        # 部门/科室名称（非独立机构）
        '法务部', '安全部', '财务室', '技术部', '运维部', '人事部',
        '行政部', '市场部', '销售部', '客服部', '研发部', '产品部',
    }
    if label in ('organization', 'ORG') and t in false_org_words:
        return True

    # City/province names mis-tagged as person names
    false_person_places = {
        '宿迁', '合肥', '兰州', '梅州', '惠州', '柳州', '漳州',
        '保定', '德阳', '绵阳', '遂宁', '衡阳', '岳阳', '德州',
        '滨州', '聊城', '临沂', '菏泽', '株洲', '湘潭', '郴州',
        '泸州', '宜宾', '遵义', '赣州', '九江', '上饶', '吉安',
        '衡水', '沧州', '廊坊', '承德', '张家口', '秦皇岛',
    }
    if label in ('private_person', 'PERSON') and t in false_person_places:
        return True

    # "修订日期：xxx" — business date prefix, not PII
    if label in ('private_date', 'DATE') and '修订' in t:
        return True

    return False


async def process_task(task_id: str, filepath: Path, filename: str, ext: str, *, skip_cache: bool = False):
    """Background processing: parse -> detect (parallel) -> store results.
    Checks file-content cache first — identical content skips detection entirely.
    Pass skip_cache=True to force re-detection (e.g. after whitelist/rules change).
    """
    async with _task_sem:
        task = tasks[task_id]
        try:
            content = filepath.read_bytes()
            fhash = _file_hash(content)
            cached = None if skip_cache else cache_get(fhash)
            if cached:
                task["status"] = "completed"
                task["progress"] = 100
                task["cached"] = True
                task["result"] = {
                    "task_id": task_id, "status": "completed", "filename": filename,
                    "cached": True,
                    **cached,
                    "completed_at": datetime.now().isoformat(),
                }
                logger.info("[%s] Cache hit (%s) – %d PII", task_id, fhash, cached.get("total_pii", 0))
                return

            task["status"] = "parsing"
            task["progress"] = 10
            logger.info("[%s] Parsing %s", task_id, filename)

            parser = PARSERS.get(ext)
            if not parser:
                raise ValueError(f"Unsupported format: {ext}")

            # Parse with dynamic timeout — scale with file size for large scanned PDFs
            file_size_mb = filepath.stat().st_size / (1024 * 1024)
            # Estimate: ~2s per MB for text extraction + ~5s per page for OCR (worst case)
            # A 40MB scanned PDF with 372 pages needs: 40*2 + 372*5 = 1940s
            # Cap at reasonable limits
            estimated_pages = max(1, int(file_size_mb * 10))  # rough: ~10 pages per MB
            PARSE_TIMEOUT = max(120, min(3600, int(file_size_mb * 2 + estimated_pages * 3)))
            logger.info("[%s] Parse timeout set to %ds (file %.1fMB, ~%d pages)",
                        task_id, PARSE_TIMEOUT, file_size_mb, estimated_pages)
            try:
                sections = await asyncio.wait_for(
                    asyncio.to_thread(parser, filepath),
                    timeout=PARSE_TIMEOUT,
                )
            except asyncio.TimeoutError:
                logger.error("[%s] Parsing timed out after %ds for %s", task_id, PARSE_TIMEOUT, filename)
                task["status"] = "error"
                task["error"] = f"文件解析超时（{PARSE_TIMEOUT}秒），可能是扫描件过大或OCR卡住"
                task["progress"] = 0
                return
            task["progress"] = 30

            # ------------------------------------------------------------------
            # Phase 1: collect ALL text segments with metadata
            # ------------------------------------------------------------------
            # Each entry: (text, section_label)
            segments: list[tuple[str, str]] = []

            # Always scan filename for PII (names, phone numbers, dates, etc.)
            name_stem = Path(filename).stem
            if name_stem.strip():
                segments.append((name_stem, "filename"))

            for sec in sections:
                if sec["type"] == "paragraph":
                    text = sec.get("text", "")
                    if text.strip():
                        segments.append((text, f"paragraph[{sec['index']}]"))
                elif sec["type"] == "table":
                    for row_idx, row in enumerate(sec["rows"]):
                        for col_idx, cell_text in enumerate(row):
                            if cell_text.strip():
                                segments.append((
                                    cell_text,
                                    f"table[{sec['index']}].row[{row_idx}].col[{col_idx}]",
                                ))

            # Pre-filter: skip very short segments (but always keep filename)
            segments = [
                s for s in segments
                if s[1] == "filename" or len(s[0]) >= MIN_SEGMENT_LENGTH
            ]

            # De-duplicate: same text only needs to be detected once
            # Map: text → (canonical_index, all_indices_with_same_text)
            text_to_indices: dict[str, list[int]] = {}
            for idx, (text, _label) in enumerate(segments):
                text_to_indices.setdefault(text, []).append(idx)
            # Unique texts to actually send to OPF (preserves order)
            unique_texts = list(dict.fromkeys(text for text, _ in segments))
            dedup_count = len(segments) - len(unique_texts)
            if dedup_count > 0:
                logger.info("[%s] De-duplicated: %d segments → %d unique texts (%d reused)",
                            task_id, len(segments), len(unique_texts), dedup_count)
            total_segments = len(segments)

            # Large document cap: skip tiny cells, limit total segments for OPF
            MAX_SEGMENTS = 500
            if total_segments > MAX_SEGMENTS:
                # Keep: filename + paragraphs + longer table cells first
                # Drop: very short table cells (< 10 chars) after the cap
                filename_seg = [s for s in segments if s[1] == "filename"]
                para_segs = [s for s in segments if s[1].startswith("paragraph")]
                table_segs = [s for s in segments if s[1].startswith("table")]

                # Sort table segments by text length desc, keep top N
                table_segs.sort(key=lambda s: len(s[0]), reverse=True)
                remaining = MAX_SEGMENTS - len(filename_seg) - len(para_segs)
                table_segs = table_segs[:max(0, remaining)]

                segments = filename_seg + para_segs + table_segs
                total_segments = len(segments)
                logger.info("[%s] Large document: capped to %d segments (dropped short cells)",
                            task_id, total_segments)

            logger.info("[%s] Detecting PII in %d segments (concurrency=%d)",
                        task_id, total_segments, OPF_CONCURRENCY)

            if total_segments == 0:
                task["parsed_sections"] = sections
                task["status"] = "completed"
                task["progress"] = 100
                task["result"] = {
                    "task_id": task_id, "status": "completed", "filename": filename,
                    "total_pii": 0, "by_type": {}, "details": [],
                    "completed_at": datetime.now().isoformat(),
                }
                logger.info("[%s] Completed – 0 PII found", task_id)
                return

            # ------------------------------------------------------------------
            # Phase 2: batch detection (de-duplicated) + post-processing
            # ------------------------------------------------------------------
            task["status"] = "detecting"
            BATCH_SIZE = 50
            DETECT_TIMEOUT = 300
            unique_results: dict[str, dict] = {}  # text → detection result

            for batch_start in range(0, len(unique_texts), BATCH_SIZE):
                batch = unique_texts[batch_start:batch_start + BATCH_SIZE]
                batch_end = min(batch_start + BATCH_SIZE, len(unique_texts))

                try:
                    batch_opf = await asyncio.wait_for(
                        detect_pii_batch(batch),
                        timeout=DETECT_TIMEOUT,
                    )
                except asyncio.TimeoutError:
                    logger.warning("[%s] Batch %d-%d timed out", task_id, batch_start, batch_end)
                    batch_opf = [{"detected_spans": []} for _ in batch]
                except Exception as e:
                    logger.warning("[%s] Batch %d-%d failed: %s", task_id, batch_start, batch_end, e)
                    batch_opf = [{"detected_spans": []} for _ in batch]

                # Post-process: dict + regex + filtering (parallel via ThreadPool)
                def _post_one(idx):
                    text = batch[idx]
                    spans_raw = batch_opf[idx].get("detected_spans", [])
                    opf_spans = [
                        {"label": s.get("label", "UNKNOWN"), "start": s.get("start", 0),
                         "end": s.get("end", 0), "text": s.get("text", "")}
                        for s in spans_raw
                    ]
                    spans = _post_process_spans(text, opf_spans)
                    redacted = text
                    for s in sorted(spans, key=lambda x: x["start"], reverse=True):
                        redacted = redacted[:s["start"]] + f"<{s['label']}>" + redacted[s["end"]:]
                    return text, {"original": text, "redacted": redacted, "spans": spans}

                from concurrent.futures import ThreadPoolExecutor
                with ThreadPoolExecutor(max_workers=4) as pool:
                    for text, result in pool.map(_post_one, range(len(batch))):
                        unique_results[text] = result

                task["progress"] = 30 + int(60 * batch_end / len(unique_texts))

                # Incremental: update partial result after each batch so frontend can show progress
                partial_details = []
                partial_pii = 0
                partial_by_type = {}
                for text, r in unique_results.items():
                    for s in r.get("spans", []):
                        partial_pii += 1
                        partial_by_type[s["label"]] = partial_by_type.get(s["label"], 0) + 1
                    if r.get("spans"):
                        partial_details.append({
                            "section": "pending",
                            "original": r["original"],
                            "redacted": r["redacted"],
                            "spans": r["spans"],
                        })
                task["result"] = {
                    "task_id": task_id, "status": "detecting",
                    "filename": filename,
                    "total_pii": partial_pii, "text_segments": total_segments,
                    "by_type": partial_by_type, "details": partial_details,
                }

                logger.info("[%s] Batch %d/%d done (%d PII so far)",
                            task_id, batch_end, len(unique_texts), partial_pii)

            # Map back to all segments (including duplicates)
            results = []
            for text, _label in segments:
                results.append(unique_results[text])

            # ------------------------------------------------------------------
            # Phase 3: assemble final details + stats
            # ------------------------------------------------------------------
            all_details: list[dict] = []
            total_pii = 0
            by_type: dict[str, int] = {}

            for idx, res in enumerate(results):
                if isinstance(res, Exception):
                    continue
                spans = res["spans"]
                if spans:
                    for s in spans:
                        total_pii += 1
                        by_type[s["label"]] = by_type.get(s["label"], 0) + 1
                    all_details.append({
                        "section": segments[idx][1],
                        "original": res["original"],
                        "redacted": res["redacted"],
                        "spans": spans,
                    })

            # Compute file stats
            all_text = "\n".join(seg[0] for seg in segments)
            char_count = len(all_text)
            if sections:
                line_count = sum(
                    len(sec.get("rows", [])) if sec["type"] == "table"
                    else sec.get("text", "").count("\n") + 1 if sec.get("text", "").strip()
                    else 0
                    for sec in sections
                )
            else:
                line_count = all_text.count("\n") + 1 if all_text else 0

            task["parsed_sections"] = sections
            task["status"] = "completed"
            task["progress"] = 100
            task["result"] = {
                "task_id": task_id,
                "status": "completed",
                "filename": filename,
                "total_pii": total_pii,
                "text_segments": total_segments,
                "char_count": char_count,
                "line_count": line_count,
                "by_type": by_type,
                "details": all_details,
                "completed_at": datetime.now().isoformat(),
            }
            # Cache for future scans
            cache_put(fhash, result)
            _save_task(task_id)
            logger.info("[%s] Completed – %d PII found (cached as %s)", task_id, total_pii, fhash)

        except Exception as exc:
            logger.exception("[%s] Error: %s", task_id, exc)
            task["status"] = "error"
            task["error"] = str(exc)
            _save_task(task_id)

    # ---------------------------------------------------------------------------
# Redaction helpers for file export
# ---------------------------------------------------------------------------

def _redact_text(text: str, details: list[dict]) -> str:
    """Apply redaction to a single text string using matched details."""
    # Build a replacement map from spans
    replacements: list[tuple[int, int, str]] = []
    for det in details:
        if det["original"] == text:
            for span in det["spans"]:
                replacements.append((span["start"], span["end"], f"<{span['label']}>"))
            break
    # Sort by start descending to avoid offset shift
    replacements.sort(key=lambda x: x[0], reverse=True)
    result = text
    for start, end, tag in replacements:
        result = result[:start] + tag + result[end:]
    return result


def export_redacted_text(filepath: Path, sections: list[dict], details: list[dict]) -> Path:
    text = sections[0]["text"] if sections else ""
    redacted = _redact_text(text, details)
    out = filepath.parent / f"redacted_{filepath.name}"
    out.write_text(redacted, encoding="utf-8")
    return out


def export_redacted_csv(filepath: Path, sections: list[dict], details: list[dict]) -> Path:
    out = filepath.parent / f"redacted_{filepath.name}"
    with open(filepath, newline="", encoding="utf-8", errors="replace") as fin, \
         open(out, "w", newline="", encoding="utf-8") as fout:
        reader = csv.reader(fin)
        writer = csv.writer(fout)
        for row in reader:
            new_row = []
            for cell in row:
                matched = [d for d in details if d["original"] == cell]
                if matched:
                    new_row.append(_redact_text(cell, matched))
                else:
                    new_row.append(cell)
            writer.writerow(new_row)
    return out


def export_redacted_xlsx(filepath: Path, sections: list[dict], details: list[dict]) -> Path:
    import openpyxl

    wb = openpyxl.load_workbook(filepath)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    matched = [d for d in details if d["original"] == cell.value]
                    if matched:
                        cell.value = _redact_text(cell.value, matched)
    out = filepath.parent / f"redacted_{filepath.name}"
    wb.save(str(out))
    wb.close()
    return out


def export_redacted_docx(filepath: Path, sections: list[dict], details: list[dict]) -> Path:
    from docx import Document

    doc = Document(str(filepath))
    for para in doc.paragraphs:
        if para.text.strip():
            matched = [d for d in details if d["original"] == para.text]
            if matched:
                new_text = _redact_text(para.text, matched)
                # Clear and re-write runs
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        matched = [d for d in details if d["original"] == para.text]
                        if matched:
                            new_text = _redact_text(para.text, matched)
                            for run in para.runs:
                                run.text = ""
                            if para.runs:
                                para.runs[0].text = new_text
    out = filepath.parent / f"redacted_{filepath.name}"
    doc.save(str(out))
    return out


def _find_cjk_font() -> str | None:
    """Find a CJK TTF/TTC font on the system."""
    import glob
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    found = glob.glob("/usr/share/fonts/**/NotoSansCJK*", recursive=True)
    return found[0] if found else None


def export_redacted_pdf(filepath: Path, sections: list[dict], details: list[dict]) -> Path:
    """Generate a PII detection report for PDF files.

    PDF cannot be reliably modified while preserving layout, so instead of
    redacting in-place, we produce a structured report listing all detected
    sensitive items with their location and type.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # Register CJK font
    font_name = "Helvetica"
    try:
        font_path = _find_cjk_font()
        if font_path:
            pdfmetrics.registerFont(TTFont("NotoCJK", font_path, subfontIndex=0))
            font_name = "NotoCJK"
        else:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
    except Exception:
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            pass

    LABEL_ZH = {
        "private_person": "姓名", "private_phone": "手机号码",
        "private_email": "电子邮箱", "account_number": "账号/证件号",
        "private_address": "地址", "private_date": "日期",
        "secret": "密码/密钥", "private_url": "URL",
        "private_bankcard": "银行卡号", "private_idcard": "身份证号",
        "organization": "机构名称", "other_person": "其他人名",
    }

    out = filepath.parent / f"pii_report_{filepath.stem}.txt"

    # Collect all PII items
    items = []
    for det in details:
        section = det.get("section", "unknown")
        for span in det.get("spans", []):
            items.append({
                "section": section,
                "type": LABEL_ZH.get(span.get("label", ""), span.get("label", "")),
                "text": span.get("text", ""),
            })

    # Build report content
    lines = []
    lines.append("=" * 60)
    lines.append("  敏感信息检测报告")
    lines.append("=" * 60)
    lines.append(f"  原始文件：{filepath.name}")
    lines.append(f"  检测时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"  检测段数：{len(sections)}")
    lines.append(f"  发现敏感项：{len(items)}")
    lines.append("=" * 60)
    lines.append("")

    if not items:
        lines.append("  ✓ 未检测到敏感信息")
        lines.append("")
    else:
        # Group by type
        by_type: dict[str, list] = {}
        for item in items:
            by_type.setdefault(item["type"], []).append(item)

        lines.append("  【按类型统计】")
        lines.append("")
        for ptype, pitems in by_type.items():
            lines.append(f"    {ptype}：{len(pitems)} 项")
        lines.append("")
        lines.append("-" * 60)
        lines.append("")
        lines.append("  【详细清单】")
        lines.append("")

        for idx, item in enumerate(items, 1):
            lines.append(f"  {idx}. [{item['type']}] {item['text']}")
            lines.append(f"     位置：{item['section']}")
            lines.append("")

    lines.append("=" * 60)
    lines.append("  报告由 OPF 隐私信息检测平台生成")
    lines.append("=" * 60)

    content = "\n".join(lines) + "\n"
    out.write_text(content, encoding="utf-8")
    return out


EXPORTERS = {
    ".txt": export_redacted_text,
    ".md": export_redacted_text,
    ".csv": export_redacted_csv,
    ".xlsx": export_redacted_xlsx,
    ".docx": export_redacted_docx,
    ".pdf": export_redacted_pdf,
}


# ---------------------------------------------------------------------------
# Task cleanup (automatic, called on list_tasks)
# ---------------------------------------------------------------------------

def _cleanup_stale_tasks() -> int:
    """Remove tasks that exceed age or count limits. Returns count removed."""
    if not tasks:
        return 0

    now = datetime.now()
    cutoff = now.timestamp() - TASK_MAX_AGE_HOURS * 3600
    removed = 0

    # Phase 1: remove by age — skip locked tasks (protected)
    stale_ids = []
    for tid, task in list(tasks.items()):
        created = task.get("created_at")
        if not created:
            continue
        # created_at can be isoformat string or unix timestamp
        if isinstance(created, str):
            try:
                ts = datetime.fromisoformat(created).timestamp()
            except (ValueError, TypeError):
                continue
        else:
            ts = created
        if ts < cutoff:
            stale_ids.append(tid)

    for tid in stale_ids:
        if tid in tasks:
            # Skip tasks that are still processing
            if tasks[tid].get("status") in ("queued", "processing", "parsing", "detecting"):
                continue
            _remove_task(tid)
            removed += 1

    # Phase 2: trim by count — keep most recent
    if len(tasks) > TASK_MAX_COUNT:
        sorted_tids = sorted(
            tasks.keys(),
            key=lambda tid: tasks[tid].get("created_at", "") or "",
            reverse=True,
        )
        for tid in sorted_tids[TASK_MAX_COUNT:]:
            if tasks[tid].get("status") in ("queued", "processing", "parsing", "detecting"):
                continue
            _remove_task(tid)
            removed += 1

    return removed


def _remove_task(task_id: str) -> None:
    """Remove a single task and its uploaded file from disk."""
    task = tasks.get(task_id)
    if not task:
        return
    filepath = task.get("filepath")
    if filepath:
        try:
            p = Path(filepath).parent
            if p.exists():
                import shutil as _shutil
                _shutil.rmtree(p, ignore_errors=True)
        except Exception:
            pass
    tasks.pop(task_id, None)


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/")
async def index():
    return FileResponse("static/index.html")


@app.get("/favicon.svg")
async def favicon():
    return FileResponse("static/favicon.svg", media_type="image/svg+xml")


@app.get("/api/tasks")
async def list_tasks():
    """List all tasks with their current status."""
    # Auto-clean stale tasks periodically (no-op most calls, does work ~1/100 times)
    _cleanup_stale_tasks()

    result = []
    for t in tasks.values():
        r = t.get("result", {})
        result.append({
            "task_id": t.get("task_id"),
            "status": t.get("status", "unknown"),
            "progress": t.get("progress", 0),
            "filename": t.get("filename"),
            "created_at": t.get("created_at"),
            "updated_at": t.get("updated_at"),
            "error": t.get("error"),
            "total_pii": r.get("total_pii", 0) if r else 0,
            "text_segments": r.get("text_segments", 0) if r else 0,
        })
    result.sort(key=lambda x: x.get("created_at") or "", reverse=True)
    return result


@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    filename = file.filename or "unknown"

    # Read file content
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, "File too large (max 50 MB)")
    if len(content) == 0:
        raise HTTPException(400, "Empty file")

    # Detect real content type via Magika (ignores extension)
    detected_ext, label, description = detect_content_type(content, filename)
    if detected_ext is None or detected_ext not in PARSERS:
        # Fallback: try using file extension if Magika misidentifies
        fallback_ext = Path(filename).suffix.lower()
        if fallback_ext and fallback_ext in PARSERS:
            logger.warning(
                "Magika detection failed for %s (label=%s), falling back to extension %s",
                filename, label, fallback_ext,
            )
            detected_ext = fallback_ext
            label = "unknown"
            description = "fallback"
        else:
            raise HTTPException(
                400,
                f"不支持的文件类型: {description} (label={label})。"
                f"当前支持: txt, md, csv, xlsx, docx, pdf"
            )
    ext = detected_ext
    logger.info("File %s detected as %s (%s)", filename, label, description)

    # Save to temp dir
    task_id = uuid.uuid4().hex[:12]
    task_dir = UPLOAD_DIR / task_id
    task_dir.mkdir(parents=True, exist_ok=True)
    filepath = task_dir / filename
    filepath.write_bytes(content)

    # Register task
    tasks[task_id] = {
        "task_id": task_id,
        "filename": filename,
        "filepath": str(filepath),
        "ext": ext,
        "detected_as": label,
        "status": "queued",
        "progress": 0,
        "created_at": datetime.now().isoformat(),
    }
    _save_task(task_id)

    # Start background processing
    asyncio.create_task(process_task(task_id, filepath, filename, ext))

    logger.info("Uploaded %s -> task %s (detected: %s)", filename, task_id, label)
    return {"task_id": task_id, "filename": filename, "detected_as": label, "status": "queued"}


@app.post("/api/batch/upload")
async def upload_batch(files: list[UploadFile] = File(...)):
    if not files:
        raise HTTPException(400, "No files provided")

    batch_id = uuid.uuid4().hex[:12]
    task_ids = []

    for file in files:
        filename = Path(file.filename or "unknown").name  # strip folder path
        content = await file.read()
        if len(content) == 0 or len(content) > MAX_FILE_SIZE:
            continue

        # Detect content type via Magika
        detected_ext, label, description = detect_content_type(content, filename)
        if detected_ext is None or detected_ext not in PARSERS:
            # Fallback: try using file extension
            fallback_ext = Path(filename).suffix.lower()
            if fallback_ext and fallback_ext in PARSERS:
                logger.warning(
                    "Magika detection failed for %s (label=%s), falling back to extension %s",
                    filename, label, fallback_ext,
                )
                detected_ext = fallback_ext
                label = "unknown"
                description = "fallback"
            else:
                continue
        ext = detected_ext

        task_id = uuid.uuid4().hex[:12]
        task_dir = UPLOAD_DIR / task_id
        task_dir.mkdir(parents=True, exist_ok=True)
        filepath = task_dir / filename
        filepath.write_bytes(content)

        tasks[task_id] = {
            "task_id": task_id,
            "filename": filename,
            "filepath": str(filepath),
            "ext": ext,
            "detected_as": label,
            "status": "queued",
            "progress": 0,
            "created_at": datetime.now().isoformat(),
        }
        _save_task(task_id)
        asyncio.create_task(process_task(task_id, filepath, filename, ext))
        task_ids.append(task_id)

    if not task_ids:
        raise HTTPException(400, "No supported files found in upload")

    batches[batch_id] = {
        "batch_id": batch_id,
        "task_ids": task_ids,
        "total": len(task_ids),
        "created_at": datetime.now().isoformat(),
    }
    logger.info("Batch %s: %d files queued", batch_id, len(task_ids))
    return {"batch_id": batch_id, "task_ids": task_ids, "total": len(task_ids)}


@app.post("/api/scan/local")
async def scan_local_path(body: dict):
    """Scan a local directory path directly — no upload needed."""
    raw_path = body.get("path", "").strip()
    if not raw_path:
        raise HTTPException(400, "path is required")

    # Map host path to container mount
    # Container mounts host home dir as /host (configurable via HOST_HOME env)
    host_home = os.getenv("HOST_HOME", "/host")
    host_path = Path(raw_path)
    home_dir = Path(os.getenv("SCAN_HOME", os.path.expanduser("~")))
    container_path = Path(host_home) / host_path.relative_to(home_dir) \
        if str(host_path).startswith(str(home_dir)) else host_path

    if not container_path.exists():
        raise HTTPException(404, f"Path not found: {raw_path}")
    if not container_path.is_dir():
        raise HTTPException(400, f"Not a directory: {raw_path}")

    batch_id = uuid.uuid4().hex[:12]
    task_ids = []

    # Walk directory and find supported files
    supported_exts = set(PARSERS.keys()) | {"jpg", "jpeg", "png", "bmp", "tiff", "tif"}
    for f in sorted(container_path.rglob("*")):
        if not f.is_file():
            continue
        ext = f.suffix.lstrip(".").lower()
        if ext not in supported_exts:
            continue
        # Skip hidden files and __MACOSX
        if any(p.startswith(".") or p == "__MACOSX" for p in f.parts):
            continue

        content = f.read_bytes()
        if len(content) == 0 or len(content) > MAX_FILE_SIZE:
            continue

        # Detect content type
        detected_ext, label, description = detect_content_type(content, f.name)
        if detected_ext is None or detected_ext not in PARSERS:
            # Fallback: try using file extension
            fallback_ext = Path(f.name).suffix.lower()
            if fallback_ext and fallback_ext in PARSERS:
                logger.warning(
                    "Magika detection failed for %s (label=%s), falling back to extension %s",
                    f.name, label, fallback_ext,
                )
                detected_ext = fallback_ext
                label = "unknown"
                description = "fallback"
            else:
                continue

        task_id = uuid.uuid4().hex[:12]
        task_dir = UPLOAD_DIR / task_id
        task_dir.mkdir(parents=True, exist_ok=True)
        copy_path = task_dir / f.name
        copy_path.write_bytes(content)

        tasks[task_id] = {
            "task_id": task_id,
            "filename": f.name,
            "filepath": str(copy_path),
            "ext": detected_ext,
            "detected_as": label,
            "status": "queued",
            "progress": 0,
            "created_at": datetime.now().isoformat(),
        }
        _save_task(task_id)
        asyncio.create_task(process_task(task_id, copy_path, f.name, detected_ext))
        task_ids.append(task_id)

    if not task_ids:
        raise HTTPException(400, f"No supported files found in {raw_path}")

    batches[batch_id] = {
        "batch_id": batch_id,
        "task_ids": task_ids,
        "total": len(task_ids),
        "created_at": datetime.now().isoformat(),
    }
    logger.info("Local scan %s: %d files from %s", batch_id, len(task_ids), raw_path)
    return {"batch_id": batch_id, "task_ids": task_ids, "total": len(task_ids)}


@app.get("/api/batch/{batch_id}/status")
async def get_batch_status(batch_id: str):
    batch = batches.get(batch_id)
    if not batch:
        raise HTTPException(404, "Batch not found")

    task_statuses = []
    for tid in batch["task_ids"]:
        t = tasks.get(tid, {})
        task_statuses.append({
            "task_id": tid,
            "filename": t.get("filename", "?"),
            "status": t.get("status", "unknown"),
            "progress": t.get("progress", 0),
            "error": t.get("error"),
            "cached": t.get("cached", False),
            "total_pii": t.get("result", {}).get("total_pii") if t.get("status") == "completed" else None,
            "by_type": t.get("result", {}).get("by_type") if t.get("status") == "completed" else None,
        })

    completed = sum(1 for t in task_statuses if t["status"] == "completed")
    errored = sum(1 for t in task_statuses if t["status"] == "error")
    total_pii = sum(t.get("total_pii") or 0 for t in task_statuses)

    return {
        "batch_id": batch_id,
        "total": batch["total"],
        "completed": completed,
        "errored": errored,
        "total_pii": total_pii,
        "status": "completed" if completed + errored >= batch["total"] else "processing",
        "tasks": task_statuses,
    }


@app.get("/api/status/{task_id}")
async def get_status(task_id: str):
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    return {
        "task_id": task_id,
        "status": task["status"],
        "progress": task.get("progress", 0),
        "filename": task["filename"],
        "error": task.get("error"),
    }


@app.get("/api/result/{task_id}")
async def get_result(task_id: str):
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task["status"] == "error":
        raise HTTPException(500, task.get("error", "Processing failed"))
    if task["status"] != "completed":
        raise HTTPException(202, detail={"status": task["status"], "progress": task.get("progress", 0)})
    return task["result"]


@app.get("/api/download/{task_id}")
async def download_redacted(task_id: str):
    task = tasks.get(task_id)
    if not task:
        raise HTTPException(404, "Task not found")
    if task["status"] != "completed":
        raise HTTPException(400, "Task not completed yet")

    filepath = Path(task["filepath"])
    ext = task["ext"]
    sections = task.get("parsed_sections", [])
    details = task["result"].get("details", [])

    exporter = EXPORTERS.get(ext)
    if not exporter:
        raise HTTPException(400, f"Export not supported for {ext}")

    try:
        out_path = exporter(filepath, sections, details)
        # PDF generates a report (.txt), not a redacted PDF
        if ext == ".pdf":
            download_name = f"pii_report_{filepath.stem}.txt"
        else:
            download_name = out_path.name
        return FileResponse(
            str(out_path),
            filename=download_name,
            media_type="application/octet-stream",
        )
    except Exception as exc:
        logger.exception("Export error for %s: %s", task_id, exc)
        raise HTTPException(500, f"Export failed: {exc}")


@app.get("/api/whitelist")
async def get_whitelist():
    return load_whitelist()


@app.post("/api/whitelist")
async def update_whitelist(data: dict):
    # Validate regex safety for all patterns (ReDoS prevention)
    for r in data.get("rules", []):
        pattern = r.get("pattern", "")
        if not pattern:
            continue
        safe, msg = is_safe_regex(pattern)
        if not safe:
            raise HTTPException(400, f"不安全的正则表达式 \"{pattern}\": {msg}")

    # Dedup rules by pattern
    seen = set()
    unique_rules = []
    for r in data.get("rules", []):
        key = r.get("pattern", "")
        if key and key not in seen:
            seen.add(key)
            unique_rules.append(r)
    data["rules"] = unique_rules
    save_whitelist(data)
    return {"status": "ok", "message": "Whitelist updated", "count": len(unique_rules)}


@app.get("/api/dictionary")
async def get_dictionary():
    return load_dictionary()


@app.post("/api/dictionary")
async def update_dictionary(data: dict):
    # Dedup words by text+type
    seen = set()
    unique_words = []
    for w in data.get("words", []):
        key = f"{w.get('text', '')}|{w.get('type', '')}"
        if key not in seen and w.get("text"):
            seen.add(key)
            unique_words.append(w)
    data["words"] = unique_words
    save_dictionary(data)
    return {"status": "ok", "message": "Dictionary updated", "count": len(unique_words)}


@app.get("/api/cache")
async def get_cache_stats():
    cache = load_scan_cache()
    return {"count": len(cache["entries"])}


@app.delete("/api/cache")
async def clear_cache():
    save_scan_cache({"version": "1.0", "entries": {}})
    tasks.clear()
    batches.clear()
    _clear_all_tasks_db()
    return {"status": "ok", "message": "Cache cleared"}


@app.post("/api/tasks/delete")
async def delete_tasks(body: dict):
    """Delete specific tasks by ID list.

    Request body: {"task_ids": ["id1", "id2", ...]}
    Returns: {"deleted": N}
    """
    task_ids = body.get("task_ids", [])
    deleted = 0
    for tid in task_ids:
        if tid in tasks:
            # Clean up uploaded file
            task = tasks[tid]
            filepath = task.get("filepath")
            if filepath:
                try:
                    import shutil as _shutil
                    p = Path(filepath).parent
                    if p.exists():
                        _shutil.rmtree(p, ignore_errors=True)
                except Exception:
                    pass
            del tasks[tid]
            _delete_task_from_db(tid)
            deleted += 1
    return {"status": "ok", "deleted": deleted}


@app.post("/api/rescan")
async def rescan_task(body: dict):
    """Re-scan a completed task using the original file.

    If the task was lost from memory (container restart), tries to
    locate the file on disk and create a new task for it.
    Request body: {"task_id": "..."}
    """
    task_id = body.get("task_id")
    if not task_id:
        raise HTTPException(400, "task_id required")

    task = tasks.get(task_id)

    # Task not in memory — try to recover from disk
    if not task:
        task_dir = UPLOAD_DIR / task_id
        if task_dir.exists():
            files = list(task_dir.iterdir())
            if files:
                filepath = files[0]
                filename = filepath.name
                ext = filepath.suffix.lower()
                tasks[task_id] = {
                    "task_id": task_id,
                    "filename": filename,
                    "filepath": str(filepath),
                    "ext": ext,
                    "status": "queued",
                    "progress": 0,
                    "result": {},
                    "error": None,
                    "created_at": datetime.now().isoformat(),
                }
                _save_task(task_id)
                task = tasks[task_id]
                asyncio.create_task(process_task(task_id, filepath, filename, ext, skip_cache=True))
                return {"status": "ok", "message": "Rescan started (recovered from disk, cache bypassed)"}
        raise HTTPException(404, "Task not found and original file missing from disk")

    filepath = Path(task["filepath"])
    if not filepath.exists():
        raise HTTPException(404, "Original file not found on disk")

    # Reset task state for re-processing
    task["status"] = "queued"
    task["progress"] = 0
    task["result"] = {}
    task["error"] = None

    filename = task["filename"]
    ext = task["ext"]

    # skip_cache=True: rescan always re-detects, never uses stale cache
    asyncio.create_task(process_task(task_id, filepath, filename, ext, skip_cache=True))
    return {"status": "ok", "message": "Rescan started (cache bypassed)"}


@app.get("/api/health/services")
async def health_services():
    """Check health of downstream services (pf_backend, OCR)."""
    services = {}

    # OPF / pf_backend (local)
    try:
        pf = _get_pf()
        services["opf"] = {"status": "ok" if pf and pf.loaded else "error"}
    except Exception:
        services["opf"] = {"status": "error"}

    # OCR (OnnxOCR)
    try:
        from onnxocr.onnx_paddleocr import ONNXPaddleOcr  # noqa: F401
        services["ocr"] = {"status": "ok"}
    except Exception:
        services["ocr"] = {"status": "unavailable"}

    return services
